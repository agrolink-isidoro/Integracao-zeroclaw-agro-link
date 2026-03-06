"""
Parser de arquivos de Máquinas (xlsx/csv/pdf/docx).

Responsabilidade: ler registros de manutenção, abastecimento e parada
e gerar drafts de Actions do tipo manutencao_maquina/abastecimento/parada_maquina.

Colunas esperadas (xlsx/csv):
  maquina_nome, maquina_id, tipo_registro, data, descricao,
  horas_trabalhadas, km_rodados, custo, tecnico, observacoes

Para PDF/DOCX: text → regex → linhas de dados.
"""

import io
import logging
import re
from datetime import datetime
from typing import Any

logger = logging.getLogger(__name__)

COLUMN_MAP = {
    "maquina": "maquina_nome", "maquina_nome": "maquina_nome",
    "equipamento": "maquina_nome", "trator": "maquina_nome",
    "nome_maquina": "maquina_nome", "machine": "maquina_nome",
    "maquina_id": "maquina_id", "id_maquina": "maquina_id", "codigo": "maquina_id",
    "tipo": "tipo_registro", "tipo_registro": "tipo_registro",
    "tipo_operacao": "tipo_registro", "operacao": "tipo_registro",
    "data": "data", "date": "data", "dt": "data",
    "descricao": "descricao", "descrição": "descricao",
    "servico": "descricao", "serviço": "descricao", "description": "descricao",
    "horas": "horas_trabalhadas", "horas_trabalhadas": "horas_trabalhadas",
    "horimetro": "horas_trabalhadas", "horímetro": "horas_trabalhadas",
    "km": "km_rodados", "km_rodados": "km_rodados", "quilometragem": "km_rodados",
    "custo": "custo", "valor": "custo", "cost": "custo", "preco": "custo",
    "tecnico": "tecnico", "técnico": "tecnico", "responsavel": "tecnico",
    "mecanico": "tecnico", "mecânico": "tecnico",
    "observacoes": "observacoes", "observações": "observacoes",
    "obs": "observacoes", "notas": "observacoes",
    # Campos específicos de abastecimento
    "litros": "quantidade_litros", "quantidade_litros": "quantidade_litros",
    "volume": "quantidade_litros", "qtd_litros": "quantidade_litros",
    "preco_litro": "valor_unitario", "valor_unitario": "valor_unitario",
    "preco_unitario": "valor_unitario", "custo_unitario": "valor_unitario",
    "produto": "produto_combustivel_nome",
    "combustivel": "produto_combustivel_nome",
    "produto_combustivel": "produto_combustivel_nome",
    "combustivel_tipo": "produto_combustivel_nome",
    "local_abastecimento": "local_abastecimento", "local": "local_abastecimento",
    # Campos específicos de manutenção
    "tipo_manutencao": "tipo_registro", "tipo_reparo": "tipo_registro",
    "prioridade": "prioridade",
    "custo_mao_obra": "custo_mao_obra",
    "status": "status",
}

TIPO_ACTION_MAP = {
    "manutencao": "manutencao_maquina",
    "manutenção": "manutencao_maquina",
    "revisao": "manutencao_maquina",
    "revisão": "manutencao_maquina",
    "reparo": "manutencao_maquina",
    "abastecimento": "abastecimento",
    "combustivel": "abastecimento",
    "combustível": "abastecimento",
    "parada": "parada_maquina",
    "parada_manutencao": "parada_maquina",
    "troca_oleo": "manutencao_maquina",
    "troca óleo": "manutencao_maquina",
}


def _normalizar_coluna(col: str) -> str:
    col = col.strip().lower().replace(" ", "_").replace("-", "_")
    return COLUMN_MAP.get(col, col)


def _parse_float(value: Any) -> float | None:
    if value is None or value == "":
        return None
    try:
        return float(str(value).replace(",", "."))
    except (ValueError, TypeError):
        return None


def _parse_data(value: Any) -> str | None:
    if not value:
        return None
    if hasattr(value, "strftime"):
        return value.strftime("%Y-%m-%d")
    s = str(value).strip()
    for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y"):
        try:
            return datetime.strptime(s, fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    return s


def _inferir_tipo_action(tipo_raw: str) -> str:
    """Mapeia tipo_registro textual para action_type canônico."""
    key = tipo_raw.strip().lower().replace(" ", "_")
    return TIPO_ACTION_MAP.get(key, "manutencao_maquina")


def _row_to_draft(row: dict) -> dict:
    tipo_raw = row.get("tipo_registro") or "manutencao"
    action_type = _inferir_tipo_action(tipo_raw)
    draft: dict = {
        "maquina_nome": row.get("maquina_nome") or "",
        "maquina_id": row.get("maquina_id") or "",
        "tipo_registro": tipo_raw,
        "data": _parse_data(row.get("data")) or "",
        "descricao": row.get("descricao") or "",
        "horas_trabalhadas": _parse_float(row.get("horas_trabalhadas")),
        "km_rodados": _parse_float(row.get("km_rodados")),
        "custo": _parse_float(row.get("custo")),
        "tecnico": row.get("tecnico") or "",
        "observacoes": row.get("observacoes") or "",
    }
    if action_type == "abastecimento":
        # Campos extras específicos de abastecimento
        draft["quantidade_litros"] = _parse_float(row.get("quantidade_litros"))
        draft["valor_unitario"] = _parse_float(row.get("valor_unitario"))
        draft["produto_combustivel_nome"] = row.get("produto_combustivel_nome") or ""
        draft["local_abastecimento"] = row.get("local_abastecimento") or ""
    elif action_type == "manutencao_maquina":
        # Campos extras específicos de manutenção
        draft["prioridade"] = row.get("prioridade") or "media"
        draft["custo_mao_obra"] = _parse_float(row.get("custo_mao_obra"))
        draft["status"] = row.get("status") or "aberta"
    return {"action_type": action_type, "draft_data": draft}


# ── xlsx ─────────────────────────────────────────────────────────────────────

def parse_xlsx(file_bytes: bytes) -> list[dict]:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []
    headers = [_normalizar_coluna(str(c) if c else "") for c in rows[0]]
    return [
        _row_to_draft(dict(zip(headers, row)))
        for row in rows[1:]
        if any(row)
    ]


# ── csv ──────────────────────────────────────────────────────────────────────

def parse_csv(file_bytes: bytes) -> list[dict]:
    import csv
    content = file_bytes.decode("utf-8-sig", errors="replace")
    reader = csv.DictReader(io.StringIO(content))
    return [
        _row_to_draft({_normalizar_coluna(k): v for k, v in row.items()})
        for row in reader
        if any(row.values())
    ]


# ── pdf ──────────────────────────────────────────────────────────────────────

def parse_pdf(file_bytes: bytes) -> list[dict]:
    """
    Extrai texto do PDF via pdfplumber e tenta encontrar linhas de dados
    usando padrões de data + maquina/tipo.
    """
    import pdfplumber

    fulltext = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            fulltext += (page.extract_text() or "") + "\n"

    return _parse_text_lines(fulltext)


def _parse_text_lines(text: str) -> list[dict]:
    """
    Regex simples: busca linhas que contêm data + padrões de manutenção.
    Formato esperado (por linha):
      DD/MM/YYYY  <maquina>  <tipo>  <descricao>  [custo]
    """
    date_pattern = re.compile(
        r"(\d{2}/\d{2}/\d{4})\s+(.+?)(?:\s{2,}|\t)(.+?)(?:\s{2,}|\t)(.+?)(?:\s+([\d.,]+))?\s*$"
    )
    drafts = []
    for line in text.splitlines():
        m = date_pattern.match(line.strip())
        if not m:
            continue
        data, maquina, tipo, descricao, custo = m.groups()
        drafts.append(
            _row_to_draft({
                "data": data,
                "maquina_nome": maquina.strip(),
                "tipo_registro": tipo.strip(),
                "descricao": descricao.strip(),
                "custo": custo,
            })
        )
    return drafts


# ── docx ─────────────────────────────────────────────────────────────────────

def parse_docx(file_bytes: bytes) -> list[dict]:
    """
    Extrai tabelas de Word (.docx) via python-docx.
    Tabelas com cabeçalho são tratadas como xlsx.
    Texto puro é tratado como PDF text.
    """
    import docx

    document = docx.Document(io.BytesIO(file_bytes))
    drafts: list[dict] = []

    for table in document.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        headers = [_normalizar_coluna(h) for h in rows[0]]
        for row_cells in rows[1:]:
            row = dict(zip(headers, row_cells))
            if any(row.values()):
                drafts.append(_row_to_draft(row))

    if not drafts:
        # Fallback: texto puro
        fulltext = "\n".join(p.text for p in document.paragraphs)
        drafts = _parse_text_lines(fulltext)

    return drafts


# ── interface principal ──────────────────────────────────────────────────────

def parse(file_bytes: bytes, mime_type: str, filename: str) -> list[dict]:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext in ("xlsx", "xls") or "spreadsheet" in mime_type:
        return parse_xlsx(file_bytes)
    if ext == "csv" or mime_type == "text/csv":
        return parse_csv(file_bytes)
    if ext == "pdf" or mime_type == "application/pdf":
        return parse_pdf(file_bytes)
    if ext == "docx" or "wordprocessingml" in mime_type:
        return parse_docx(file_bytes)

    raise ValueError(f"Formato não suportado para Máquinas: {ext or mime_type}")
