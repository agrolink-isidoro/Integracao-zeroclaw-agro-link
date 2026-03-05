"""
WebSocket Consumer para o chat com Isidoro.

Rota: ws://backend/ws/chat/

Protocolo de mensagens (JSON):
  Cliente → Servidor:
    { "type": "message", "text": "..." }
    { "type": "upload_result", "upload_id": "uuid", "module": "estoque" }
    { "type": "ping" }

  Servidor → Cliente:
    { "type": "message", "text": "...", "sender": "isidoro", "timestamp": "..." }
    { "type": "typing", "is_typing": true }
    { "type": "action_created", "action_id": "uuid", "module": "...", "action_type": "..." }
    { "type": "error", "message": "..." }
    { "type": "pong" }

O consumer autentica via JWT no query string ou header:
  ws://backend/ws/chat/?token=<jwt>
"""

import json
import logging
import os
from datetime import datetime

from channels.generic.websocket import AsyncWebsocketConsumer
from channels.db import database_sync_to_async

logger = logging.getLogger(__name__)

# Limite de caracteres do conteúdo do arquivo injetado no contexto do LLM
_FILE_CONTENT_LIMIT = 12_000


def _extract_text_from_file(path: str, mime_type: str, filename: str) -> str:
    """
    Extrai texto legível de um arquivo para ser injetado no contexto do LLM.

    Suporta: XML, CSV, TXT, JSON, markdown, PDF, Excel, Word (.docx/.doc), OpenOffice (.odt/.ods/.odp).
    """
    ext = os.path.splitext(filename)[1].lower()

    # ── texto simples ──────────────────────────────────────────────────────
    if ext in ('.xml', '.csv', '.txt', '.json', '.md', '.yaml', '.yml', '.html'):
        try:
            with open(path, 'r', encoding='utf-8', errors='replace') as fh:
                content = fh.read(_FILE_CONTENT_LIMIT + 500)
            if len(content) > _FILE_CONTENT_LIMIT:
                content = content[:_FILE_CONTENT_LIMIT] + '\n…[conteúdo truncado]'
            return content
        except OSError as e:
            return f"[Não foi possível ler o arquivo: {e}]"

    # ── PDF ───────────────────────────────────────────────────────────────
    if ext == '.pdf' or 'pdf' in mime_type:
        try:
            import pypdf  # pypdf >= 3.x
            text_parts = []
            with open(path, 'rb') as fh:
                reader = pypdf.PdfReader(fh)
                for page in reader.pages:
                    text_parts.append(page.extract_text() or '')
            content = '\n'.join(text_parts)
            if len(content) > _FILE_CONTENT_LIMIT:
                content = content[:_FILE_CONTENT_LIMIT] + '\n…[conteúdo truncado]'
            return content or '[PDF sem texto extraível]'
        except ImportError:
            pass
        try:
            import PyPDF2
            text_parts = []
            with open(path, 'rb') as fh:
                reader = PyPDF2.PdfReader(fh)
                for page in reader.pages:
                    text_parts.append(page.extract_text() or '')
            content = '\n'.join(text_parts)
            if len(content) > _FILE_CONTENT_LIMIT:
                content = content[:_FILE_CONTENT_LIMIT] + '\n…[conteúdo truncado]'
            return content or '[PDF sem texto extraível]'
        except (ImportError, Exception) as e:
            return f"[PDF recebido, mas não foi possível extrair o texto: {e}]"

    # ── Excel ─────────────────────────────────────────────────────────────
    if ext in ('.xlsx', '.xls'):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            lines = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                lines.append(f'=== Aba: {sheet} ===')
                for row in ws.iter_rows(values_only=True):
                    row_str = '\t'.join('' if v is None else str(v) for v in row)
                    if row_str.strip():
                        lines.append(row_str)
            content = '\n'.join(lines)
            if len(content) > _FILE_CONTENT_LIMIT:
                content = content[:_FILE_CONTENT_LIMIT] + '\n…[conteúdo truncado]'
            return content
        except (ImportError, Exception) as e:
            return f"[Excel recebido, mas não foi possível extrair: {e}]"

    # ── Word .docx ────────────────────────────────────────────────────────
    if ext == '.docx':
        try:
            import docx  # python-docx
            doc = docx.Document(path)
            lines = [para.text for para in doc.paragraphs if para.text.strip()]
            # Também extrai tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        lines.append(row_text)
            content = '\n'.join(lines)
            if len(content) > _FILE_CONTENT_LIMIT:
                content = content[:_FILE_CONTENT_LIMIT] + '\n…[conteúdo truncado]'
            return content or '[.docx sem texto extraível]'
        except ImportError:
            return '[.docx recebido, mas python-docx não está instalado no servidor]'
        except Exception as e:
            return f'[.docx recebido, mas não foi possível extrair: {e}]'

    # ── Word .doc (formato legado) ─────────────────────────────────────────
    if ext == '.doc':
        try:
            import subprocess
            result = subprocess.run(
                ['antiword', path], capture_output=True, text=True, timeout=15
            )
            content = result.stdout
            if content.strip():
                if len(content) > _FILE_CONTENT_LIMIT:
                    content = content[:_FILE_CONTENT_LIMIT] + '\n…[conteúdo truncado]'
                return content
        except (FileNotFoundError, Exception):
            pass
        return '[.doc (formato legado) recebido — converta para .docx para melhor suporte]'

    # ── OpenOffice Writer .odt ────────────────────────────────────────────
    if ext == '.odt':
        try:
            from odf.opendocument import load as odf_load
            from odf.text import P
            from odf.teletype import extractText
            doc = odf_load(path)
            lines = []
            for para in doc.getElementsByType(P):
                text = extractText(para).strip()
                if text:
                    lines.append(text)
            content = '\n'.join(lines)
            if len(content) > _FILE_CONTENT_LIMIT:
                content = content[:_FILE_CONTENT_LIMIT] + '\n…[conteúdo truncado]'
            return content or '[.odt sem texto extraível]'
        except ImportError:
            return '[.odt recebido, mas odfpy não está instalado no servidor]'
        except Exception as e:
            return f'[.odt recebido, mas não foi possível extrair: {e}]'

    # ── OpenOffice Calc .ods ──────────────────────────────────────────────
    if ext == '.ods':
        try:
            from odf.opendocument import load as odf_load
            from odf.table import Table, TableRow, TableCell
            from odf.teletype import extractText
            doc = odf_load(path)
            lines = []
            for sheet in doc.getElementsByType(Table):
                lines.append(f'=== Aba: {sheet.getAttribute("name")} ===')
                for row in sheet.getElementsByType(TableRow):
                    cells = row.getElementsByType(TableCell)
                    row_text = '\t'.join(extractText(c).strip() for c in cells)
                    if row_text.strip():
                        lines.append(row_text)
            content = '\n'.join(lines)
            if len(content) > _FILE_CONTENT_LIMIT:
                content = content[:_FILE_CONTENT_LIMIT] + '\n…[conteúdo truncado]'
            return content or '[.ods sem dados extraíveis]'
        except ImportError:
            return '[.ods recebido, mas odfpy não está instalado no servidor]'
        except Exception as e:
            return f'[.ods recebido, mas não foi possível extrair: {e}]'

    # ── OpenOffice Impress .odp ───────────────────────────────────────────
    if ext == '.odp':
        try:
            from odf.opendocument import load as odf_load
            from odf.text import P
            from odf.teletype import extractText
            doc = odf_load(path)
            lines = []
            for para in doc.getElementsByType(P):
                text = extractText(para).strip()
                if text:
                    lines.append(text)
            content = '\n'.join(lines)
            if len(content) > _FILE_CONTENT_LIMIT:
                content = content[:_FILE_CONTENT_LIMIT] + '\n…[conteúdo truncado]'
            return content or '[.odp sem texto extraível]'
        except ImportError:
            return '[.odp recebido, mas odfpy não está instalado no servidor]'
        except Exception as e:
            return f'[.odp recebido, mas não foi possível extrair: {e}]'

    # ── fallback ──────────────────────────────────────────────────────────
    return f"[Arquivo '{filename}' recebido (tipo: {mime_type}). Conteúdo não pôde ser lido como texto.]"


class IsidoroChatConsumer(AsyncWebsocketConsumer):
    """
    Consumer WebSocket para o chat com o agente Isidoro.

    Cada conexão WebSocket representa uma sessão de chat de um usuário.
    O agente Isidoro é instanciado por tenant (compartilhado no layer de channels).
    """

    async def connect(self):
        """Aceita conexão, autentica, e inicia sessão Isidoro."""
        user = self.scope.get("user")
        if not user or not user.is_authenticated:
            await self.close(code=4001)
            return

        self.user = user
        self.tenant = self.scope.get("tenant")

        # Tenant is optional — superusers (admin) may not have one
        self.tenant_id = str(self.tenant.id) if self.tenant else "global"
        self.user_id = str(user.id)
        self.group_name = f"chat_{self.tenant_id}_{self.user_id}"

        # Entra no grupo de channels (para broadcast de eventos)
        await self.channel_layer.group_add(self.group_name, self.channel_name)
        await self.accept()

        # Inicia o agente Isidoro (com fallback se API key ausente)
        try:
            self.isidoro = self._get_isidoro_agent()
        except (ValueError, Exception) as exc:
            logger.error("Falha ao inicializar Isidoro agent: %s", exc)
            self.isidoro = None

        logger.info("WebSocket chat connected: user=%s tenant=%s", self.user_id, self.tenant_id)

        # Briefing contextualizado do dia (chama ferramentas: safras, pendências, estoque, máquinas)
        if self.isidoro is None:
            await self._send_error(
                "⚠️ Isidoro está offline — a chave de API do LLM não está configurada. "
                "Configure `ISIDORO_API_KEY` no `.env` e reinicie o backend."
            )
            return

        await self._send_typing(True)
        try:
            tenant_nome = getattr(self.tenant, "nome", "Sua Fazenda")
            greeting = await self.isidoro.initialize_session(
                tenant_id=self.tenant_id,
                user_id=self.user_id,
                tenant_nome=tenant_nome,
            )
            await self._send_typing(False)
            await self._send_message(text=greeting.text, sender="isidoro")
        except Exception as exc:
            logger.exception("Erro ao gerar briefing inicial: %s", exc)
            await self._send_typing(False)
            await self._send_message(
                text="Olá! Sou o Isidoro, seu assistente agrícola. Como posso ajudar hoje?",
                sender="isidoro",
            )

    async def disconnect(self, close_code):
        """Sai do grupo ao desconectar."""
        if hasattr(self, "group_name"):
            await self.channel_layer.group_discard(self.group_name, self.channel_name)
        logger.info("WebSocket chat disconnected: user=%s code=%s", getattr(self, "user_id", "?"), close_code)

    async def receive(self, text_data):
        """Recebe mensagem do cliente e despacha para o handler correto."""
        try:
            data = json.loads(text_data)
        except json.JSONDecodeError:
            await self._send_error("Mensagem inválida (JSON malformado)")
            return

        msg_type = data.get("type", "message")

        if msg_type == "ping":
            await self.send(json.dumps({"type": "pong"}))
        elif msg_type == "message":
            await self._handle_chat_message(data.get("text", "").strip())
        elif msg_type == "upload_result":
            await self._handle_upload_notification(data)
        else:
            await self._send_error(f"Tipo de mensagem não reconhecido: {msg_type}")

    # ─────────────────────────────────────────────────────────────────────────

    async def _handle_chat_message(self, text: str):
        """Processa mensagem de texto do usuário."""
        if not text:
            return

        if self.isidoro is None:
            await self._send_error(
                "Isidoro está offline — chave de API não configurada. "
                "Configure ISIDORO_API_KEY no .env e reinicie o backend."
            )
            return

        # Mostra "digitando..."
        await self._send_typing(True)

        try:
            tenant_nome = getattr(self.tenant, "nome", "Sua Fazenda")
            response = await self.isidoro.chat(
                user_message=text,
                tenant_id=self.tenant_id,
                user_id=self.user_id,
                tenant_nome=tenant_nome,
            )

            await self._send_typing(False)
            await self._send_message(text=response.text, sender="isidoro")

            if response.error:
                logger.warning("Isidoro error: %s", response.error)

        except Exception as exc:
            logger.exception("Erro no _handle_chat_message: %s", exc)
            await self._send_typing(False)
            await self._send_error("Erro interno ao processar sua mensagem.")

    async def _handle_upload_notification(self, data: dict):
        """Lê o arquivo, injeta no histórico do Isidoro e pede análise inicial."""
        upload_id = data.get("upload_id", "")
        module = data.get("module", "")
        filename = data.get("filename", "arquivo")

        await self._send_typing(True)

        # Busca o registro do upload no banco para obter o caminho do arquivo
        upload_obj = await self._get_upload(upload_id)

        file_content: str
        if upload_obj and upload_obj.caminho_arquivo and os.path.exists(upload_obj.caminho_arquivo):
            file_content = _extract_text_from_file(
                path=upload_obj.caminho_arquivo,
                mime_type=upload_obj.mime_type or "",
                filename=upload_obj.nome_original or filename,
            )
        else:
            file_content = f"[Arquivo '{filename}' registrado mas o conteúdo não está disponível no servidor.]"

        # Injeta o conteúdo no histórico da sessão Isidoro
        tenant_nome = getattr(self.tenant, "nome", "Sua Fazenda")
        self.isidoro.inject_file_context(
            tenant_id=self.tenant_id,
            user_id=self.user_id,
            filename=upload_obj.nome_original if upload_obj else filename,
            content=file_content,
            tenant_nome=tenant_nome,
        )

        # Pede ao Isidoro para apresentar os dados principais do arquivo
        prompt = (
            f"O usuário acabou de enviar o arquivo '{filename}' (módulo: {module}). "
            "Apresente de forma resumida e clara os dados principais que encontrou no conteúdo "
            "e pergunte o que ele gostaria de fazer (ex: registrar no sistema, consultar algo, etc.)."
        )

        try:
            response = await self.isidoro.chat(
                user_message=prompt,
                tenant_id=self.tenant_id,
                user_id=self.user_id,
                tenant_nome=tenant_nome,
            )
            await self._send_typing(False)
            await self._send_message(text=response.text, sender="isidoro")
        except Exception as exc:
            logger.exception("Erro ao processar upload no Isidoro: %s", exc)
            await self._send_typing(False)
            await self._send_message(
                text=(
                    f"Recebi o arquivo '{filename}'!\n\n"
                    "O que você gostaria de fazer com ele?\n"
                    "Por exemplo: registrar entrada de estoque, lançar nota fiscal, "
                    "importar operações agrícolas, analisar manutenções…"
                ),
                sender="isidoro",
            )

        logger.info("Upload processed: upload_id=%s module=%s filename=%s content_chars=%d",
                    upload_id, module, filename, len(file_content))

    @database_sync_to_async
    def _get_upload(self, upload_id: str):
        """Busca o UploadedFile no banco de dados."""
        if not upload_id:
            return None
        try:
            from apps.actions.models import UploadedFile
            return UploadedFile.objects.get(id=upload_id)
        except Exception:
            return None

    # ─────────────────────────────────────────────────────────────────────────

    async def _send_message(self, text: str, sender: str = "isidoro"):
        """Envia mensagem de texto ao cliente."""
        await self.send(json.dumps({
            "type": "message",
            "text": text,
            "sender": sender,
            "timestamp": datetime.utcnow().isoformat() + "Z",
        }))

    async def _send_typing(self, is_typing: bool):
        """Envia indicador de 'digitando...'."""
        await self.send(json.dumps({"type": "typing", "is_typing": is_typing}))

    async def _send_error(self, message: str):
        """Envia mensagem de erro ao cliente."""
        await self.send(json.dumps({"type": "error", "message": message}))

    def _get_isidoro_agent(self):
        """Cria ou reutiliza o agente Isidoro para este tenant."""
        from django.conf import settings
        from zeroclaw_tools.integrations.agrolink import IsidoroAgent

        return IsidoroAgent(
            base_url=getattr(settings, "AGROLINK_API_URL", "http://localhost:8000/api"),
            jwt_token=self._get_isidoro_jwt(),
            model=getattr(settings, "ISIDORO_LLM_MODEL", "glm-5"),
            api_key=getattr(settings, "ISIDORO_API_KEY", None),
            llm_base_url=getattr(settings, "ISIDORO_LLM_BASE_URL", None),
        )

    def _get_isidoro_jwt(self) -> str:
        """Obtém o JWT do Isidoro para chamadas à API Agrolink."""
        from django.conf import settings
        return getattr(settings, "ISIDORO_JWT_TOKEN", "")

    # ─────────────────────────────────────────────────────────────────────────
    # Channel layer group handlers (para eventos server-side → cliente)

    async def chat_message(self, event):
        """Handler para mensagens enviadas pelo channel layer (ex: Action aprovada)."""
        await self.send(json.dumps(event.get("data", {})))
